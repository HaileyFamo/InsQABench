########## 实现文档问答的框架 ##########
# 此处使用了GPT API为例子


import time
import os
import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import ChatOpenAI
from langchain_community.embeddings import HuggingFaceBgeEmbeddings
from langchain_core.documents import Document
from langchain_core.messages import AIMessage, HumanMessage
from langchain_core.runnables import RunnablePassthrough, RunnableBranch
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.schema import Document
import docx
from PyPDF2 import PdfReader
from streamlit_pdf_viewer import pdf_viewer


import os

api_key = os.getenv('OPENAI_API_KEY')
if api_key is None:
    raise ValueError("OPENAI_API_KEY environment variable is not set.")


os.environ["TOKENIZERS_PARALLELISM"] = "false"


st.set_page_config(page_title="InsLLM，与您的保险合同对话", page_icon=":speech_balloon:", layout="wide")

st.markdown(
    """<style>
.chat-message {
    padding: 1.5rem; border-radius: 0.5rem; margin-bottom: 1rem; display: flex
}
.chat-message.user {
    background-color: #2b313e
}
.chat-message.bot {
    background-color: #475063
}
.chat-message .avatar {
  width: 20%;
}
.chat-message .avatar img {
  max-width: 78px;
  max-height: 78px;
  border-radius: 50%;
  object-fit: cover;
}
.chat-message .message {
  width: 80%;
  padding: 0 1.5rem;
  color: #fff;
}
.stDeployButton {
            visibility: hidden;
        }
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}

.block-container {
    padding: 2rem 4rem 2rem 4rem;
}

.st-emotion-cache-16txtl3 {
    padding: 3rem 1.5rem;
}
</style>
<script>
function scrollToBottom() {
    var chatContainer = document.querySelector('.block-container');
    chatContainer.scrollTop = chatContainer.scrollHeight;
}
</script>
# """,
    unsafe_allow_html=True,
)


# bot_template = """
# <div class="chat-message bot">
#     <div class="avatar">
#         <img src="https://img.icons8.com/?size=100&id=79UfeEN6JkZ8&format=png&color=000000" style="max-height: 78px; max-width: 78px; border-radius: 50%; object-fit: cover;">
#     </div>
#     <div class="message">{{MSG}}</div>
# </div>
# """

# user_template = """
# <div class="chat-message user">
#     <div class="avatar">
#         <img src="https://img.icons8.com/?size=100&id=23301&format=png&color=000000" >
#     </div>    
#     <div class="message">{{MSG}}</div>
# </div>
# """

bot_template = """
<div class="chat-message bot">
    <div class="avatar">
        <img src="https://img.icons8.com/?size=100&id=79UfeEN6JkZ8&format=png&color=000000" style="max-height: 50px; max-width: 50px; border-radius: 50%; object-fit: cover;">
    </div>
    <div class="message">{{MSG}}</div>
</div>
"""

user_template = """
<div class="chat-message user">
    <div class="avatar">
        <img src="https://img.icons8.com/?size=100&id=23301&format=png&color=000000" style="max-height: 50px; max-width: 50px;">
    </div>    
    <div class="message">{{MSG}}</div>
</div>
"""

def get_pdf_text(pdf_docs):

    docs = []
    for document in pdf_docs:
        if document.type == "application/pdf":
            pdf_reader = PdfReader(document)
            for idx, page in enumerate(pdf_reader.pages):
                docs.append(
                    Document(
                        page_content=page.extract_text(),
                        metadata={"source": f"{document.name} on page {idx}"},
                    )
                )
        elif (
            document.type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            doc = docx.Document(document)
            for idx, paragraph in enumerate(doc.paragraphs):
                docs.append(
                    Document(
                        page_content=paragraph.text,
                        metadata={"source": f"{document.name} in paragraph {idx}"},
                    )
                )
        elif document.type == "text/plain":
            text = document.getvalue().decode("utf-8")
            docs.append(Document(page_content=text, metadata={"source": document.name}))

    return docs


def get_text_chunks(docs):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=0)

    docs_chunks = text_splitter.split_documents(docs)
    return docs_chunks


def get_vectorstore(docs_chunks):
    embeddings = HuggingFaceBgeEmbeddings(model_name="BAAI/bge-base-zh-v1.5")
    vectorstore = FAISS.from_documents(docs_chunks, embedding=embeddings)
    return vectorstore


example = """
"input": "你是一个保险行业的从业者，需要为保险的消费者解答问题。下面根据保险合同的片段给出回答：\n这段文字来自的合同是：TaiPing公众责任保险（2019版）\n片段内容：第三条 在保险期间内，被保险人在本保险单明细表中列明的地点范围内依法从事生产、经营等活动时，由于意外事故造成的下列损失或费用，依照中华人民共和国法律（不包括港澳台地区法律）应由被保险人承担的经济赔偿责任，保险人按照本保险合同约定负责赔偿： （一）第三者人身伤亡或财产损失；  （二）事先经保险人书面同意的诉讼或仲裁费用（以下简称“法律费用”）； （三）发生保险责任事故后，被保险人为缩小或减少对第三者人身伤亡或财产损失的赔偿责任所支付的必要的、合理的费用。  保险人对上述（一）（二）（三）项保险责任的每次事故赔偿总金额不超过本保险单明细表中列明的每次事故赔偿限额；对每次事故承担的法律费用的赔偿金额不超过每次事故赔偿限额的 10%，但合同另有约定的除外。在保险期限内，保险人对被保险人的累计赔偿总金额不得超过本保险单明细表中列明的累计赔偿限额。  责任免除 \n用户提问：保险人对每次事故的赔偿总金额有什么限制？", "output": "保险人对每次事故的赔偿总金额不超过本保险单 **明细表** 中列明的 **每次事故赔偿限额**；对每次事故承担的法律费用的赔偿金额不超过每次事故赔偿限额的 **10%**，但合同另有约定的除外。\n**专有名词解释**:\n* **明细表**：保险合同中附带的表格，详细列明保险的具体内容，如保险金额、赔偿限额、免除责任等。\n* **每次事故赔偿限额**：保险人对每次保险事故的最高赔偿金额。"
"""

def get_conversation_chain(vectorstore):
    query_transform_prompt = ChatPromptTemplate.from_messages(
        [
            MessagesPlaceholder(variable_name="messages"),
            (
                "user",
                "当用户输入一个提问时，将其改写成为更加适合检索器检索的查询，不要添加任何额外的信息，只返回你修改后的查询",
            ),
        ]
    )

    system_prompt = (
        "你是一个中国的保险专家。 "
        "你需要根据提供的保险合同中的有关片段，回答用户的问题。"
        "如果涉及到保险合同，请以合同中的条款为准。"
        "如果有些信息合同中没有，请根据常识回答。"
        "回答完毕后，对用户的问题和你的回答中出现的保险的专有名词进行解释。具体的格式是：**专有名词**： 专有名词的解释。\n"
        "使用专业而礼貌的语气回答用户的问题。"
        "下面是一个例子："
        f"{example}"
        "用户提问："
        "<context>"
        "{context}"
        "</context>"
    )   

    llm = ChatOpenAI(model="gpt-4o-mini")
    retriever = vectorstore.as_retriever()

    qa_prompt = ChatPromptTemplate.from_messages(
        [
            ("system", system_prompt),
            MessagesPlaceholder(variable_name="messages"),
        ]
    )

    query_transforming_retriever_chain = RunnableBranch(
        (
            lambda x: len(x.get("messages", [])) == 1,
            # If only one message, then we just pass that message's content to retriever
            (lambda x: x["messages"][-1].content) | retriever,
        ),
        # If messages, then we pass inputs to LLM chain to transform the query, then pass to retriever
        query_transform_prompt | llm | StrOutputParser() | retriever,
    ).with_config(run_name="chat_retriever_chain")


    document_chain = create_stuff_documents_chain(llm, qa_prompt)
    conversational_retrieval_chain = RunnablePassthrough.assign(
        context=query_transforming_retriever_chain,
    ).assign(
        answer=document_chain,
    )

    return conversational_retrieval_chain


def handle_userinput_pdf(user_question):
    # 先打印用户的问题
    st.write(
        user_template.replace("{{MSG}}", user_question),
        unsafe_allow_html=True,
    )

    chat_history = st.session_state.chat_history
    messages = []

    # 将历史对话转换为 HumanMessage 和 AIMessage 对象
    for role, content in chat_history:
        if role == "user":
            messages.append(HumanMessage(content=content))
        elif role == "assistant":
            messages.append(AIMessage(content=content))

    # 添加当前用户输入
    messages.append(HumanMessage(content=user_question))
    display_on = False
    full_content = st.empty()
    display_string = ""
    sources = []

    # 使用stream_data函数实现流式输出
    def stream_data():
        nonlocal display_on, sources
        response = st.session_state.conversation.stream(
            {
                "messages": messages,
            }
        )
        for chunk in response:
            if 'context' in chunk and len(chunk["context"]) > 0:
                sources = chunk["context"]
                display_on = True
            if display_on:
                yield chunk.get('answer', '')
                time.sleep(0.02)

    # 使用st.write_stream实现流式输出
    for chunk in stream_data():
        display_string += chunk
        # full_content.markdown(bot_template.replace("{{MSG}}", display_string), unsafe_allow_html=True)
        full_content.markdown(bot_template.replace("{{MSG}}", display_string), unsafe_allow_html=True)
        st.markdown("<script>scrollToBottom();</script>", unsafe_allow_html=True)

    # 处理sources
    if display_on:
        source_names = set([i.metadata["source"] for i in sources])
        src = "\n\n".join(source_names)
        src = f"\n\n> 来源 : \n{src}"
        display_string += src
        # full_content.markdown(bot_template.replace("{{MSG}}", display_string), unsafe_allow_html=True)
        full_content.markdown(bot_template.replace("{{MSG}}", display_string), unsafe_allow_html=True)
        st.markdown("<script>scrollToBottom();</script>", unsafe_allow_html=True)

    # 更新 chat_history
    st.session_state.chat_history.append(("user", user_question))
    st.session_state.chat_history.append(("assistant", display_string))


def show_history():
    chat_history = st.session_state.chat_history

    for i, message in enumerate(chat_history):
        if i % 2 == 0:
            st.write(
                user_template.replace("{{MSG}}", message[1]),
                unsafe_allow_html=True,
            )
        else:
            st.write(
                bot_template.replace("{{MSG}}", message[1]), unsafe_allow_html=True
            )

def main():
    # st.header("InsLLM，与您的保险合同对话。")
    st.header(":violet[_InsLLM_] ，与您的保险合同对话。")
    # 初始化会话状态
    if "conversation" not in st.session_state:
        st.session_state.conversation = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    with st.sidebar:
        st.title("文档管理")
        pdf_docs = st.file_uploader(
            "选择文件，支持pdf、txt、doc、docx格式",
            type=["pdf", "txt", "doc", "docx"],
            accept_multiple_files=True,
        )
        if st.button(
            "处理您的文档",
            on_click=lambda: setattr(st.session_state, "last_action", "pdf"),
            use_container_width=True,
        ):
            if pdf_docs:
                with st.spinner("Processing"):
                    docs = get_pdf_text(pdf_docs)
                    docs_chunks = get_text_chunks(docs)
                    vectorstore = get_vectorstore(docs_chunks)
                    st.session_state.conversation = get_conversation_chain(vectorstore)
                    st.session_state.pdf_docs = pdf_docs
            else:
                st.warning("请先上传文件")

        def clear_history():
            st.session_state.chat_history = []

        if st.session_state.chat_history:
            st.button("清空对话", on_click=clear_history, use_container_width=True)

        if "pdf_docs" in st.session_state:
            for pdf_doc in st.session_state.pdf_docs:
                st.write(f"**{pdf_doc.name}**")
                pdf_viewer(pdf_doc.getvalue())

    with st.container():
        user_question = st.chat_input("请输入您的问题")

    with st.container(height=500):
        show_history()
        if user_question:
            if st.session_state.conversation is not None:
                handle_userinput_pdf(user_question)
            else:
                st.warning("请先上传文件")

if __name__ == "__main__":
    main()